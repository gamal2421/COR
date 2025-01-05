from faker import Faker
from docx import Document

fake = Faker()

def generate_random_cv():
    # Generate random personal information
    name = fake.name()
    email = fake.email()
    phone = fake.phone_number()
    education = fake.job()  # You can customize this for more specific fields
    experience = fake.text(max_nb_chars=200)

    # Create a Word document
    doc = Document()
    doc.add_heading(f"Resume: {name}", 0)
    
    doc.add_paragraph(f"Name: {name}")
    doc.add_paragraph(f"Email: {email}")
    doc.add_paragraph(f"Phone: {phone}")
    doc.add_paragraph(f"Education: {education}")
    doc.add_paragraph(f"Experience: {experience}")

    # Save to a file
    file_name = f"{name}_CV.docx"
    doc.save(file_name)
    print(f"Random CV generated: {file_name}")


for _ in range(20):  # Change 10 to however many CVs you need
    generate_random_cv()
